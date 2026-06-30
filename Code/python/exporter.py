#!/usr/bin/env python3
"""
exporter.py — PDF & DOCX generation for Pipeline Author

Generates properly paginated, formatted documents from the pipeline stage data.
Used by server.py via POST /api/export/pdf and /api/export/docx.
"""

import json
import os
import tempfile
from io import BytesIO

# ─── PDF (fpdf2) ───
try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

# ─── DOCX (python-docx) ───
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    Document = None

# ─── Helpers ───

def _safe(s, default=""):
    """Return string value, never None."""
    return (s or "").strip()


def _yesno(val):
    """Convert yes/no value to display string."""
    if val == "yes":
        return "✅ Yes"
    elif val == "no":
        return "❌ No"
    return ""


# ─── Stage 1 data extraction ───

def _extract_stage1(sd, section_defs):
    """Extract Stage 1 PRD data into a structured dict for document generation."""
    inputs = sd.get("inputs", {})
    functions = []
    for i, name in enumerate(sd.get("functionNames", [])):
        if name and name.strip():
            functions.append({
                "number": i + 1,
                "name": name,
                "summary": (sd.get("functionSummaries") or [""] * (i + 1))[i] or "",
                "scope": (sd.get("functionScoping") or [[]] * (i + 1))[i] or [],
            })
    ext_counts = sd.get("externalCounts", {})
    bff_products = []
    for i in range(1, (ext_counts.get("bff", 0) or 0) + 1):
        v = inputs.get(f"D3.3.bff_{i}", "")
        if v:
            bff_products.append(v)
    db_products = []
    for i in range(1, (ext_counts.get("perm", 0) or 0) + 1):
        v = inputs.get(f"D3.4.perm_{i}", "")
        if v:
            db_products.append(v)
    imm_products = []
    for i in range(1, (ext_counts.get("imm", 0) or 0) + 1):
        v = inputs.get(f"D3.5.imm_{i}", "")
        if v:
            imm_products.append(v)

    infrastructure = {}
    infra_items = [
        ("D1.6.1.1", "Public Webapp", ["D1.6.1.2", "D1.6.1.3"]),
        ("D1.6.2.1", "Private Webapp (Admin)", ["D1.6.2.2", "D1.6.2.3"]),
        ("D1.6.3.1", "Public Android App", ["D1.6.3.2", "D1.6.3.3"]),
        ("D1.6.4.1", "Private Android App (Admin)", ["D1.6.4.2", "D1.6.4.3"]),
        ("D1.6.5.1", "Public BFF", ["D1.6.5.2", "D1.6.5.3"]),
        ("D1.6.6.1", "Private BFF", ["D1.6.6.2", "D1.6.6.3"]),
        ("D1.6.7.1", "Permanent Database", ["D1.6.7.2", "D1.6.7.3"]),
        ("D1.6.8.1", "Permanent DB Functions", ["D1.6.8.2", "D1.6.8.3"]),
        ("D1.6.9.1", "In-memory / Cache Database", ["D1.6.9.2", "D1.6.9.3"]),
        ("D1.6.10.1", "In-memory DB Functions", ["D1.6.10.2", "D1.6.10.3"]),
    ]
    for key, label, followups in infra_items:
        val = inputs.get(key, "")
        if val == "yes":
            details = {}
            for f_id in followups:
                details[f_id] = _safe(inputs.get(f_id, ""))
            infrastructure[key] = {"label": label, "details": details}

    return {
        "productName": _safe(inputs.get("D1.1", "")),
        "businessPurpose": _safe(inputs.get("D1.2.1", "")),
        "newUserWorkflow": _safe(inputs.get("D1.2.2", "")),
        "userTypes": {
            "readOnly": inputs.get("D1.2.3.1") == "yes",
            "readOnlyDesc": _safe(inputs.get("D1.2.3.1a", "")),
            "writeOnly": inputs.get("D1.2.3.2") == "yes",
            "writeOnlyDesc": _safe(inputs.get("D1.2.3.2a", "")),
            "premium": inputs.get("D1.2.3.3") == "yes",
            "premiumFeatures": _safe(inputs.get("D1.2.3.3a", "")),
            "premiumSubTypes": int(inputs.get("D1.2.3.4", "0") or "0"),
            "adminPage": inputs.get("D1.2.3.5") == "yes",
            "superAdminPage": inputs.get("D1.2.3.6") == "yes",
        },
        "github": inputs.get("D1.3") == "yes",
        "functionCount": sd.get("functionCount", 0),
        "functions": functions,
        "infrastructure": infrastructure,
        "externalLinkages": {
            "hasExternal": inputs.get("D3.1") == "yes",
            "interfaces": inputs.get("D3.2", []),
            "bffProducts": bff_products,
            "databaseProducts": db_products,
            "inMemoryProducts": imm_products,
        },
        "d5Results": sd.get("d5Results", ""),
        "d4ContextDiagram": sd.get("d4ContextDiagram", ""),
    }


# ─── PDF Generation ───

class PipelinePDF(FPDF):
    """Custom PDF class with header/footer for pipeline exports."""

    def __init__(self, title="Pipeline Export"):
        super().__init__()
        self.pipeline_title = title
        self.set_auto_page_break(auto=True, margin=20)

    def header(self):
        if self.page_no() > 1:
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(120, 120, 120)
            self.cell(
                0, 8, self.pipeline_title, align="L"
            )
            self.cell(
                0, 8, f"Page {self.page_no()}", align="R", new_x="LMARGIN", new_y="NEXT"
            )
            self.line(10, 14, 200, 14)
            self.ln(4)

    def footer(self):
        if self.page_no() > 1:
            self.set_y(-15)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(160, 160, 160)
            self.cell(0, 10, f"Generated by Pipeline Author", align="C")


def _pdf_add_section_title(pdf, title):
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(30, 27, 58)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_draw_color(99, 102, 241)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)


def _pdf_add_subsection(pdf, title):
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(75, 70, 112)
    pdf.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)


def _pdf_add_label_value(pdf, label, value, indent=0):
    x = 10 + indent
    pdf.set_x(x)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(60, 60, 60)
    pdf.cell(50, 6, label, align="L")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(20, 20, 20)
    if value:
        # Handle multi-line values
        for line in value.split("\n"):
            pdf.set_x(x + 52)
            pdf.multi_cell(130, 5, line)
    else:
        pdf.cell(0, 6, "-")
    pdf.ln(1)


def _pdf_add_body(pdf, text, indent=0):
    pdf.set_x(10 + indent)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(20, 20, 20)
    for line in text.split("\n"):
        pdf.set_x(10 + indent)
        pdf.multi_cell(180 - indent, 5, line.strip() or " ")
    pdf.ln(2)


def _pdf_add_table_row(pdf, cells, col_widths, bold=False, fill=False):
    """Add a single row to a table with given column widths."""
    pdf.set_font("Helvetica", "B" if bold else "", 8)
    if fill:
        pdf.set_fill_color(240, 240, 255)
    for i, cell in enumerate(cells):
        pdf.cell(col_widths[i] if i < len(col_widths) else 30, 7, str(cell), border=1, fill=fill)
    pdf.ln()


def generate_pdf_bytes(stage_data, pipeline_def, title="Pipeline Export"):
    """
    Generate a PDF as bytes from pipeline stage data.
    
    Args:
        stage_data: dict of stageId -> stage data dict (from frontend stageData)
        pipeline_def: list of stage definition dicts (from PIPELINE array on frontend)
        title: document title
    
    Returns:
        bytes: PDF file content
    """
    if FPDF is None:
        raise RuntimeError("fpdf2 is not installed. Run: pip install fpdf2")

    pdf = PipelinePDF(title=title)
    pdf.add_page()

    # ── Cover / Title page ──
    pdf.set_font("Helvetica", "B", 24)
    pdf.set_text_color(30, 27, 58)
    pdf.ln(40)
    pdf.cell(0, 15, "Pipeline Author", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 14)
    pdf.set_text_color(99, 102, 241)
    pdf.cell(0, 10, "Full Pipeline Export", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 10)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 8, f"Generated: {title.replace('Pipeline Export - ', '')}", align="C", new_x="LMARGIN", new_y="NEXT")
    total_stages = len(pipeline_def)
    completed = sum(1 for s in pipeline_def if stage_data.get(str(s["id"]) if isinstance(s.get("id"), int) else s.get("id"), {}).get("completed"))
    pdf.cell(0, 8, f"Stages: {completed}/{total_stages} completed", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.add_page()

    # ── Iterate through stages ──
    for stage_idx, stage in enumerate(pipeline_def):
        sid = str(stage.get("id", stage_idx + 1))
        sd = stage_data.get(sid, {})
        is_stage1 = stage.get("isStage1PRD", False)
        stage_name = stage.get("name", f"Stage {sid}")

        # Stage header
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_text_color(30, 27, 58)
        pdf.cell(0, 12, f"Stage {sid}: {stage_name}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(30, 27, 58)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)

        # Stage type badge
        stage_type = stage.get("type", "")
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 6, f"Type: {stage_type}  |  Models: {', '.join(stage.get('models', []))}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

        # ── Stage 1 PRD: detailed structured output ──
        if is_stage1:
            s1 = _extract_stage1(sd, None)
            _pdf_add_stage1_sections(pdf, s1)
        else:
            # ── Stages 2-9: standard manual + AI deliverables ──
            manual = stage.get("manualDeliverables", [])
            ai = stage.get("aiDeliverables", [])

            # Manual inputs
            if manual:
                _pdf_add_section_title(pdf, "Manual Inputs")
                for d in manual:
                    did = d.get("id", "")
                    label = d.get("label", did)
                    val = sd.get("manualInputs", {}).get(did, "")
                    _pdf_add_subsection(pdf, label)
                    if val and val.strip():
                        _pdf_add_body(pdf, val)
                    else:
                        pdf.set_font("Helvetica", "I", 8)
                        pdf.set_text_color(180, 180, 180)
                        pdf.cell(0, 6, "(Not provided)", new_x="LMARGIN", new_y="NEXT")
                        pdf.ln(2)

            # AI outputs
            if ai:
                _pdf_add_section_title(pdf, "AI-Generated Outputs")
                for d in ai:
                    did = d.get("id", "")
                    label = d.get("label", did)
                    val = sd.get("aiOutputs", {}).get(did, "")
                    _pdf_add_subsection(pdf, label)
                    if val and val.strip():
                        _pdf_add_body(pdf, val)
                    else:
                        pdf.set_font("Helvetica", "I", 8)
                        pdf.set_text_color(180, 180, 180)
                        pdf.cell(0, 6, "(Not generated yet)", new_x="LMARGIN", new_y="NEXT")
                        pdf.ln(2)

            # Gate reviews
            gate_reviews = stage.get("gateReviews", [])
            if gate_reviews:
                _pdf_add_section_title(pdf, "Gate Review Responses")
                for r in gate_reviews:
                    rid = r.get("id", "")
                    question = r.get("question", "")
                    notes = sd.get("reviewNotes", {}).get(rid, "")
                    _pdf_add_subsection(pdf, rid)
                    _pdf_add_body(pdf, f"Q: {question}")
                    if notes and notes.strip():
                        _pdf_add_body(pdf, f"A: {notes}")
                    pdf.ln(2)

        # Completed status
        if sd.get("completed"):
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(16, 185, 129)
            pdf.cell(0, 6, "✓ Stage completed", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(6)

        # Page break between stages
        if stage_idx < len(pipeline_def) - 1:
            pdf.add_page()

    # Write to bytes
    return pdf.output()


def _pdf_add_stage1_sections(pdf, s1):
    """Add Stage 1 PRD sections to the PDF."""
    # ── Product Basics ──
    _pdf_add_section_title(pdf, "📦 Product Basics")
    _pdf_add_label_value(pdf, "Product Name", s1["productName"])
    _pdf_add_label_value(pdf, "Business Purpose", s1["businessPurpose"])
    _pdf_add_label_value(pdf, "New User Workflow", s1["newUserWorkflow"])

    # User types
    _pdf_add_subsection(pdf, "User Types")
    ut = s1["userTypes"]
    _pdf_add_label_value(pdf, "Read-Only Users", "✅ Yes" if ut["readOnly"] else "❌ No")
    if ut["readOnly"] and ut["readOnlyDesc"]:
        _pdf_add_label_value(pdf, "  Description", ut["readOnlyDesc"])
    _pdf_add_label_value(pdf, "Write-Only Users", "✅ Yes" if ut["writeOnly"] else "❌ No")
    if ut["writeOnly"] and ut["writeOnlyDesc"]:
        _pdf_add_label_value(pdf, "  Description", ut["writeOnlyDesc"])
    _pdf_add_label_value(pdf, "Premium Users", "✅ Yes" if ut["premium"] else "❌ No")
    if ut["premium"] and ut["premiumFeatures"]:
        _pdf_add_label_value(pdf, "  Premium Features", ut["premiumFeatures"])
    _pdf_add_label_value(pdf, "Premium Sub-Types", str(ut["premiumSubTypes"]))
    _pdf_add_label_value(pdf, "Admin Page", "✅ Yes" if ut["adminPage"] else "❌ No")
    _pdf_add_label_value(pdf, "Super-Admin Page", "✅ Yes" if ut["superAdminPage"] else "❌ No")
    pdf.ln(2)

    # GitHub
    _pdf_add_section_title(pdf, "📁 Repository Setup")
    _pdf_add_label_value(pdf, "GitHub Repository", "✅ Yes" if s1["github"] else "❌ No (Local only)")
    pdf.ln(2)

    # ── Infrastructure ──
    _pdf_add_section_title(pdf, "🖥️ Infrastructure")
    infra = s1["infrastructure"]
    if infra:
        for key, item in infra.items():
            _pdf_add_subsection(pdf, item["label"])
            for f_id, f_val in item["details"].items():
                if f_val:
                    _pdf_add_label_value(pdf, "", f_val)
    else:
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(180, 180, 180)
        pdf.cell(0, 6, "(No infrastructure items selected)", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # ── Functions ──
    _pdf_add_section_title(pdf, "⚙️ Functions")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(60, 60, 60)
    pdf.cell(0, 6, f"Total functions: {s1['functionCount']}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    for fn in s1["functions"]:
        _pdf_add_subsection(pdf, f"Function {fn['number']}: {fn['name']}")
        pdf.set_x(15)
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_text_color(60, 60, 60)
        pdf.cell(30, 5, "Summary:", align="L")
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(20, 20, 20)
        pdf.multi_cell(145, 5, fn["summary"] or "(Not provided)")
        if fn.get("scope"):
            pdf.set_x(15)
            pdf.set_font("Helvetica", "B", 8)
            pdf.set_text_color(60, 60, 60)
            pdf.cell(30, 5, "Scope:", align="L")
            pdf.set_font("Helvetica", "", 8)
            pdf.set_text_color(20, 20, 20)
            pdf.cell(0, 5, ", ".join(fn["scope"]), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # ── External Linkages ──
    _pdf_add_section_title(pdf, "🔗 External Linkages")
    ext = s1["externalLinkages"]
    _pdf_add_label_value(pdf, "Has External Products", "✅ Yes" if ext["hasExternal"] else "❌ No")
    if ext.get("interfaces"):
        _pdf_add_label_value(pdf, "Interfaces", ", ".join(ext["interfaces"]))
    if ext.get("bffProducts"):
        _pdf_add_label_value(pdf, "BFF Products", ", ".join(ext["bffProducts"]))
    if ext.get("databaseProducts"):
        _pdf_add_label_value(pdf, "Database Products", ", ".join(ext["databaseProducts"]))
    if ext.get("inMemoryProducts"):
        _pdf_add_label_value(pdf, "In-Memory Products", ", ".join(ext["inMemoryProducts"]))
    pdf.ln(2)

    # ── D5 Auto-Checks ──
    if s1.get("d5Results"):
        _pdf_add_section_title(pdf, "🔍 Auto-Generated Checks (D5)")
        _pdf_add_body(pdf, s1["d5Results"])
        pdf.ln(2)

    # ── D4 Context Diagram ──
    if s1.get("d4ContextDiagram"):
        _pdf_add_section_title(pdf, "📐 C4 Context Diagram (D4)")
        _pdf_add_body(pdf, s1["d4ContextDiagram"])
        pdf.ln(2)


# ─── DOCX Generation ───

def _docx_add_heading(doc, text, level=1):
    """Add a heading styled like the document."""
    h = doc.add_heading(text, level=level)
    return h


def _docx_add_table(doc, headers, rows, col_widths=None):
    """Add a styled table to the document."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def _docx_add_section(doc, title):
    _docx_add_heading(doc, title, level=2)


def _docx_add_subsection(doc, title):
    _docx_add_heading(doc, title, level=3)


def _docx_add_body(doc, text):
    if not text or not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text.strip())
    run.font.size = Pt(10)
    run.font.name = "Calibri"


def _docx_add_label(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_value = p.add_run(str(value or "-"))
    run_value.font.size = Pt(10)


def _docx_add_stage1(doc, s1):
    """Add Stage 1 PRD content to the DOCX."""
    _docx_add_section(doc, "Product Basics")
    _docx_add_label(doc, "Product Name", s1["productName"])
    _docx_add_label(doc, "Business Purpose", s1["businessPurpose"])
    _docx_add_label(doc, "New User Workflow", s1["newUserWorkflow"])

    _docx_add_subsection(doc, "User Types")
    ut = s1["userTypes"]
    _docx_add_label(doc, "Read-Only Users", "Yes" if ut["readOnly"] else "No")
    if ut["readOnly"] and ut["readOnlyDesc"]:
        _docx_add_label(doc, "  Description", ut["readOnlyDesc"])
    _docx_add_label(doc, "Write-Only Users", "Yes" if ut["writeOnly"] else "No")
    if ut["writeOnly"] and ut["writeOnlyDesc"]:
        _docx_add_label(doc, "  Description", ut["writeOnlyDesc"])
    _docx_add_label(doc, "Premium Users", "Yes" if ut["premium"] else "No")
    if ut["premium"] and ut["premiumFeatures"]:
        _docx_add_label(doc, "  Premium Features", ut["premiumFeatures"])
    _docx_add_label(doc, "Premium Sub-Types", str(ut["premiumSubTypes"]))
    _docx_add_label(doc, "Admin Page", "Yes" if ut["adminPage"] else "No")
    _docx_add_label(doc, "Super-Admin Page", "Yes" if ut["superAdminPage"] else "No")

    _docx_add_section(doc, "Repository Setup")
    _docx_add_label(doc, "GitHub Repository", "Yes" if s1["github"] else "No (Local only)")

    _docx_add_section(doc, "Infrastructure")
    infra = s1["infrastructure"]
    if infra:
        for key, item in infra.items():
            _docx_add_subsection(doc, item["label"])
            for f_id, f_val in item["details"].items():
                if f_val:
                    _docx_add_body(doc, f_val)
    else:
        _docx_add_body(doc, "(No infrastructure items selected)")

    _docx_add_section(doc, "Functions")
    _docx_add_body(doc, f"Total functions: {s1['functionCount']}")
    for fn in s1["functions"]:
        _docx_add_subsection(doc, f"Function {fn['number']}: {fn['name']}")
        _docx_add_label(doc, "Summary", fn["summary"] or "(Not provided)")
        if fn.get("scope"):
            _docx_add_label(doc, "Scope", ", ".join(fn["scope"]))

    _docx_add_section(doc, "External Linkages")
    ext = s1["externalLinkages"]
    _docx_add_label(doc, "Has External Products", "Yes" if ext["hasExternal"] else "No")
    if ext.get("interfaces"):
        _docx_add_label(doc, "Interfaces", ", ".join(ext["interfaces"]))
    if ext.get("bffProducts"):
        _docx_add_label(doc, "BFF Products", ", ".join(ext["bffProducts"]))
    if ext.get("databaseProducts"):
        _docx_add_label(doc, "Database Products", ", ".join(ext["databaseProducts"]))
    if ext.get("inMemoryProducts"):
        _docx_add_label(doc, "In-Memory Products", ", ".join(ext["inMemoryProducts"]))

    if s1.get("d5Results"):
        _docx_add_section(doc, "Auto-Generated Checks (D5)")
        _docx_add_body(doc, s1["d5Results"])

    if s1.get("d4ContextDiagram"):
        _docx_add_section(doc, "C4 Context Diagram (D4)")
        _docx_add_body(doc, s1["d4ContextDiagram"])


def generate_docx_bytes(stage_data, pipeline_def, title="Pipeline Export"):
    """
    Generate a DOCX file as bytes from pipeline stage data.
    
    Args:
        stage_data: dict of stageId -> stage data dict
        pipeline_def: list of stage definition dicts (from PIPELINE on frontend)
        title: document title
    
    Returns:
        bytes: DOCX file content
    """
    if Document is None:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()

    # ── Document styles ──
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ── Title ──
    _docx_add_heading(doc, "Pipeline Author - Full Export", level=0)

    # Subtitle
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(99, 102, 241)

    total_stages = len(pipeline_def)
    completed = sum(1 for s in pipeline_def if stage_data.get(str(s["id"]) if isinstance(s.get("id"), int) else s.get("id"), {}).get("completed"))
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Stages: {completed}/{total_stages} completed")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ── Iterate through stages ──
    for stage_idx, stage in enumerate(pipeline_def):
        sid = str(stage.get("id", stage_idx + 1))
        sd = stage_data.get(sid, {})
        is_stage1 = stage.get("isStage1PRD", False)
        stage_name = stage.get("name", f"Stage {sid}")

        # Stage heading
        _docx_add_heading(doc, f"Stage {sid}: {stage_name}", level=1)

        # Stage metadata
        p = doc.add_paragraph()
        run = p.add_run(f"Type: {stage.get('type', '')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)
        run = p.add_run(f"  |  Models: {', '.join(stage.get('models', []))}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)

        if is_stage1:
            s1 = _extract_stage1(sd, None)
            _docx_add_stage1(doc, s1)
        else:
            manual = stage.get("manualDeliverables", [])
            ai = stage.get("aiDeliverables", [])

            if manual:
                _docx_add_section(doc, "Manual Inputs")
                for d in manual:
                    did = d.get("id", "")
                    label = d.get("label", did)
                    val = sd.get("manualInputs", {}).get(did, "")
                    _docx_add_subsection(doc, label)
                    if val and val.strip():
                        _docx_add_body(doc, val)
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run("(Not provided)")
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(180, 180, 180)
                        run.italic = True

            if ai:
                _docx_add_section(doc, "AI-Generated Outputs")
                for d in ai:
                    did = d.get("id", "")
                    label = d.get("label", did)
                    val = sd.get("aiOutputs", {}).get(did, "")
                    _docx_add_subsection(doc, label)
                    if val and val.strip():
                        _docx_add_body(doc, val)
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run("(Not generated yet)")
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(180, 180, 180)
                        run.italic = True

            gate_reviews = stage.get("gateReviews", [])
            if gate_reviews:
                _docx_add_section(doc, "Gate Review Responses")
                for r in gate_reviews:
                    rid = r.get("id", "")
                    question = r.get("question", "")
                    notes = sd.get("reviewNotes", {}).get(rid, "")
                    _docx_add_subsection(doc, rid)
                    _docx_add_body(doc, f"Q: {question}")
                    if notes and notes.strip():
                        _docx_add_body(doc, f"A: {notes}")

        # Completed status
        if sd.get("completed"):
            p = doc.add_paragraph()
            run = p.add_run("✓ Stage completed")
            run.bold = True
            run.font.color.rgb = RGBColor(16, 185, 129)

        # Page break between stages
        if stage_idx < len(pipeline_def) - 1:
            doc.add_page_break()

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()